from openpyxl import load_workbook
from docx import Document


# File paths
excel_file_path = 'files/example.xlsx'
word_template_path = 'files/template.docx'
output_word_path_template = 'files/Risultato_part_{}.docx'

# Helper function to replace placeholders in a paragraph's runs
def replace_placeholder_in_paragraph(paragraph, placeholders):
    for run in paragraph.runs:
        for placeholder, value in placeholders.items():
            if placeholder in run.text:
                print(f"Replacing '{placeholder}' with '{value}' in: {run.text}")
                run.text = run.text.replace(placeholder, value)

# Helper function to clear specific table cells
def clear_remaining_cells(table, start_index):
    """
    Clears specific table cells starting from the calculated index based on written rows.
    Target cells are [1, 3, 4, 6, 7, 9, 10, 12] (1-based), translated to [0, 2, 3, 5, 6, 8, 9, 11] (0-based).
    """
    target_indices = [0, 2, 3, 5, 6, 8, 9, 11]  # 0-based indices for target cells
    remaining_indices = target_indices[start_index:]
    for index in remaining_indices:
        cell = table._cells[index]
        for paragraph in cell.paragraphs:
            print(f"Clearing remaining cell at 1-based index {index + 1}: {paragraph.text}")
            paragraph.text = ""

# Function to check if the output file is locked
def ensure_file_writable(file_path):
    try:
        # Attempt to open the file in write mode
        with open(file_path, 'a'):
            pass
    except PermissionError:
        print(f"Output file '{file_path}' is currently open. Please close it and press Enter to continue.")
        input()  # Wait for user to close the file
        ensure_file_writable(file_path)  # Retry after user intervention

# Load the Excel workbook
wb = load_workbook(excel_file_path)
sheet = wb.active

# Convert Excel rows to a list (excluding the header row)
excel_rows = list(sheet.iter_rows(min_row=2, values_only=True))  # Skip header row

# Split data into chunks of 8 rows
chunk_size = 8
chunks = [excel_rows[i:i + chunk_size] for i in range(0, len(excel_rows), chunk_size)]

# Process each chunk and create a separate Word document
for chunk_index, chunk in enumerate(chunks, start=1):
    # Load a fresh template for each chunk
    doc = Document(word_template_path)
    
    # Retrieve the first table from the Word document
    table = doc.tables[0]
    
    # Specific target cell indices to populate
    target_indices = [0, 2, 3, 5, 6, 8, 9, 11]  # 0-based indices corresponding to 1, 3, 4, 6, 7, 9, 10, 12
    
    # Check if the table can accommodate the chunk
    if len(chunk) > len(target_indices):
        print("The Word table does not have enough non-blank cells to accommodate the chunk data.")
        print(f"Add more cells or reduce the data to {len(target_indices)} items.")
        continue
    
    # Populate the template with data from the chunk
    for row_index, excel_row in enumerate(chunk):
        cell = table._cells[target_indices[row_index]]  # Get the specific target cell
        
        # Unpack Excel row values
        marca, modello, descrizione, codice, dettaglio, prezzo_listino, prezzo_scontato = excel_row
        
        # Format decimal numbers and handle empty or None values
        prezzo_listino_formatted = (
            "" if prezzo_listino is None else f"{prezzo_listino:,.2f}".replace(".", ",")
        )
        prezzo_scontato_formatted = (
            "" if prezzo_scontato is None else f"{prezzo_scontato:,.2f}".replace(".", ",")
        )
        
        # Handle empty dettaglio and remove parentheses if needed
        dettaglio_formatted = f"({str(dettaglio).strip()})" if dettaglio else ""
        
        # Prepare replacement values
        placeholders = {
            "[Marca]": str(marca).strip() if marca else "",
            "[Modello]": str(modello).strip() if modello else "",
            "[Descrizione]": str(descrizione).strip() if descrizione else "",
            "[Codice]": str(codice).strip() if codice else "",
            "([Dettaglio])": dettaglio_formatted,
            "[Prezzo Listino]": prezzo_listino_formatted,
            "[Prezzo Scontato]": prezzo_scontato_formatted,
        }
        
        # Replace placeholders in the cell's paragraphs while retaining formatting
        for paragraph in cell.paragraphs:
            replace_placeholder_in_paragraph(paragraph, placeholders)
    
    # Clear remaining cells in the table
    clear_remaining_cells(table, len(chunk))
    
    # Determine output file path
    output_word_path = output_word_path_template.format(chunk_index)
    
    # Check if the file is writable
    ensure_file_writable(output_word_path)
    
    # Save the updated Word document for the current chunk
    doc.save(output_word_path)
    print(f"Document saved as {output_word_path}")


